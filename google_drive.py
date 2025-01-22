from flask import Flask, redirect, url_for, request, session, render_template, flash, send_file, jsonify
from google_auth_oauthlib.flow import Flow
from googleapiclient.discovery import build
import json
import os
import io
from googleapiclient.http import MediaIoBaseDownload, MediaFileUpload  # Add this import
from google.oauth2 import service_account
from docx import Document
from fuzzywuzzy import fuzz
from flask_wtf.csrf import CSRFProtect  # Add this import

# Allow insecure transport for development
os.environ['OAUTHLIB_INSECURE_TRANSPORT'] = '1'

CONFIG_FILE_PATH = 'c:/Users/phill/Desktop/PIRG Site/bin/PIRG-Site/config.json'  # Update with the actual path to the config file

app = Flask(__name__)
csrf = CSRFProtect(app)  # Initialize CSRFProtect

def google_drive_auth():
    with open(CONFIG_FILE_PATH, 'r') as f:
        config = json.load(f)
    flow = Flow.from_client_config(
        config,
        scopes=config['web']['scopes'],  # Read scopes from config.json
        redirect_uri='http://localhost:5000/google_drive_callback'  # Ensure this matches the Google Cloud Console
    )
    authorization_url, state = flow.authorization_url(access_type='offline', include_granted_scopes='true', prompt='consent')
    session['state'] = state
    return redirect(authorization_url)

def google_drive_callback():
    state = session['state']
    with open(CONFIG_FILE_PATH, 'r') as f:
        config = json.load(f)
    flow = Flow.from_client_config(
        config,
        scopes=config['web']['scopes'],  # Read scopes from config.json
        state=state,
        redirect_uri='http://localhost:5000/google_drive_callback'  # Ensure this matches the Google Cloud Console
    )
    flow.fetch_token(authorization_response=request.url)
    credentials = flow.credentials
    session['credentials'] = credentials_to_dict(credentials)
    return redirect(url_for('picker'))

def picker_route():
    with open(CONFIG_FILE_PATH, 'r') as f:
        config = json.load(f)
    return render_template('admin_pages/picker.html', developer_key=config['web']['developer_key'], client_id=config['web']['client_id'], scopes=config['web']['scopes'])

def credentials_to_dict(credentials):
    return {
        'token': credentials.token,
        'refresh_token': credentials.refresh_token,
        'token_uri': credentials.token_uri,
        'client_id': credentials.client_id,
        'client_secret': credentials.client_secret,
        'scopes': credentials.scopes
    }

def select_folder():
    folder_id = request.form['folder_id']
    with open(CONFIG_FILE_PATH, 'r+') as f:
        config = json.load(f)
        config['folder_id'] = folder_id
        f.seek(0)
        json.dump(config, f, indent=4)
        f.truncate()
    return redirect(url_for('doc_db'))

def search_google_drive(query, config_file=CONFIG_FILE_PATH):
    """
    Search for files in a Google Drive folder that match the query in titles and contents.
    
    :param query: The query string to search for in file names and contents.
    :param config_file: Path to the config file containing Google Drive credentials and folder ID.
    :return: A list of matching file contents.
    """
    with open(config_file, 'r') as f:
        config = json.load(f)
    
    credentials_file = config['credentials_file']
    folder_id = config['folder_id']
    
    SCOPES = ['https://www.googleapis.com/auth/drive.readonly']
    creds = service_account.Credentials.from_service_account_file(credentials_file, scopes=SCOPES)
    service = build('drive', 'v3', credentials=creds)

    results = service.files().list(q=f"'{folder_id}' in parents", pageSize=1000, fields="files(id, name, mimeType)").execute()
    items = results.get('files', [])

    matching_files = []
    for item in items:
        file_id = item['id']
        file_name = item['name']
        if fuzz.partial_ratio(query.lower(), file_name.lower()) > 70:
            matching_files.append(file_name)
        else:
            if item['mimeType'] == 'application/vnd.google-apps.document':
                request = service.files().export_media(fileId=file_id, mimeType='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                fh = io.BytesIO()
                downloader = MediaIoBaseDownload(fh, request)
                done = False
                while done is False:
                    status, done = downloader.next_chunk()
                fh.seek(0)
                doc = Document(fh)
                full_text = []
                for para in doc.paragraphs:
                    full_text.append(para.text)
                if fuzz.partial_ratio(query.lower(), '\n'.join(full_text).lower()) > 70:
                    matching_files.append(file_name)
            else:
                request = service.files().get_media(fileId=file_id)
                fh = io.BytesIO()
                downloader = MediaIoBaseDownload(fh, request)
                done = False
                while done is False:
                    status, done = downloader.next_chunk()
                fh.seek(0)
                content = fh.read().decode('utf-8', errors='ignore')
                if fuzz.partial_ratio(query.lower(), content.lower()) > 70:
                    matching_files.append(file_name)
    return matching_files

def sync_google_drive(config_file=CONFIG_FILE_PATH):
    """
    Sync files from a Google Drive folder to the database.
    
    :param config_file: Path to the config file containing Google Drive credentials and folder ID.
    """
    from models import DocumentModel  # Import here to avoid circular import

    with open(config_file, 'r') as f:
        config = json.load(f)
    
    credentials_file = config['credentials_file']
    folder_id = config['folder_id']
    
    SCOPES = ['https://www.googleapis.com/auth/drive.readonly']
    creds = service_account.Credentials.from_service_account_file(credentials_file, scopes=SCOPES)
    service = build('drive', 'v3', credentials=creds)

    results = service.files().list(q=f"'{folder_id}' in parents", pageSize=1000, fields="files(id, name, mimeType)").execute()
    items = results.get('files', [])

    for item in items:
        file_id = item['id']
        file_name = item['name']
        request = service.files().get_media(fileId=file_id)
        fh = io.BytesIO()
        downloader = MediaIoBaseDownload(fh, request)
        done = False
        while done is False:
            status, done = downloader.next_chunk()
        fh.seek(0)
        if item['mimeType'] == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            doc = Document(fh)
            full_text = '\n'.join([para.text for para in doc.paragraphs])
        else:
            full_text = fh.read().decode('utf-8', errors='ignore')
        DocumentModel.add_document(file_name, full_text)

def sync_folder(folder_id):
    """
    Sync the local 'documents' folder with the specified Google Drive folder.
    
    :param folder_id: The ID of the Google Drive folder to sync with.
    """
    with open(CONFIG_FILE_PATH, 'r') as f:
        config = json.load(f)
    
    credentials_file = config['credentials_file']
    
    SCOPES = ['https://www.googleapis.com/auth/drive']
    creds = service_account.Credentials.from_service_account_file(credentials_file, scopes=SCOPES)
    service = build('drive', 'v3', credentials=creds)

    local_folder = './documents'  # Path to the local folder to sync
    for root, dirs, files in os.walk(local_folder):
        for file_name in files:
            file_path = os.path.join(root, file_name)
            file_metadata = {
                'name': file_name,
                'parents': [folder_id]
            }
            media = MediaFileUpload(file_path, resumable=True)
            service.files().create(body=file_metadata, media_body=media, fields='id').execute()

@app.route('/sync_folder', methods=['POST'])
@csrf.exempt
def sync_folder_route():
    data = request.get_json()
    folder_id = data.get('folder_id')
    if not folder_id:
        return jsonify({'success': False, 'error': 'No folder ID provided'}), 400
    try:
        sync_folder(folder_id)
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/login')
def login():
    return google_drive_auth()

@app.route('/google_drive_callback')
def google_drive_callback_route():
    return google_drive_callback()

@app.route('/select_folder', methods=['POST'])
def select_folder_route():
    return select_folder()

if __name__ == '__main__': 
    app.run(debug=True)